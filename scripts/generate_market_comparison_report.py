from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
OUT_DIR.mkdir(parents=True, exist_ok=True)
MD_PATH = OUT_DIR / "market_comparison_2026-04-08.md"
DOCX_PATH = OUT_DIR / "market_comparison_2026-04-08.docx"
PDF_PATH = OUT_DIR / "market_comparison_2026-04-08.pdf"

TITLE = "Сравнение Rdegon SIEM с enterprise SIEM/SOAR платформами мирового и российского рынка"
SUBTITLE = "Исправленная и расширенная версия, актуально на 8 апреля 2026 года"

EXEC_SUMMARY = [
    "Предыдущая версия отчёта была неполной по vendor coverage и была испорчена кодировкой при генерации через inline Python в PowerShell. Эта версия собирается из UTF-8 скрипта и расширяет baseline по мировому и российскому рынку.",
    "Наша платформа уже выглядит как рабочий SOC/SIEM stack: есть раздельные ingest, processing, storage и web-слои, Kafka, ClickHouse, stream/batch correlation, incident management, cases, entities, assets, connectors, vulnerability management, threat intel, response actions, OIDC/Vault/governance, OpenClaw AI и Telegram bot.",
    "По архитектурной прозрачности и управляемости система сильнее типичного самописного SOC-стенда, но по масштабу, глубине контента, количеству готовых интеграций и зрелости packaged enterprise-функций пока уступает лидерам мирового и российского рынка.",
    "Главный разрыв с мировыми лидерами: scale-out storage/query architecture, зрелая UEBA/entity analytics, сертифицированный integration/content ecosystem, зрелая multi-tenancy/MSSP-модель, облачные/XDR/identity fusion сценарии и industrialized content operations.",
    "Главный разрыв с сильными российскими решениями: готовые локализованные контент-паки, регуляторные и отраслевые сценарии, зрелые IRP/SOAR-процессы, внедренческая экосистема и развитый российский пакет интеграций.",
]

MARKET_SCOPE = {
    "Мировой рынок": [
        "Splunk Enterprise Security (+ Splunk SOAR)",
        "Microsoft Sentinel",
        "Google Security Operations",
        "Palo Alto Cortex XSIAM",
        "IBM QRadar SIEM / QRadar Suite",
        "OpenText ArcSight Enterprise Security Manager",
        "Rapid7 InsightIDR / Incident Command",
        "Securonix Unified Defense SIEM",
        "Exabeam Security Operations Platform",
        "Elastic Security / Elastic SIEM",
    ],
    "Российский рынок": [
        "Positive Technologies MaxPatrol SIEM",
        "Kaspersky SIEM / KUMA",
        "R-Vision SIEM / SOAR",
        "Security Vision SIEM / SOAR / IRP",
        "UserGate SIEM",
        "Solar SIEM",
        "RuSIEM",
    ],
}

COMPARISON_ROWS = [
    {
        "criterion": "Архитектура и масштабирование",
        "ours": "Есть раздельные ingest, processing, storage и web-слои, Kafka, ClickHouse, hot/cold split и часть HA-механик. Но live-storage пока не полноценный распределённый shard/replica-кластер.",
        "global": "У Splunk, Sentinel, Google SecOps, XSIAM, QRadar, ArcSight, Rapid7, Securonix, Exabeam и Elastic зрелее deployment-модели, масштабирование, облачный или гибридный operating model и более формализованный HA/DR.",
        "russian": "Сильные российские платформы заметно лучше подготовлены к крупным on-prem/high-load внедрениям и централизованному корпоративному SOC, а Solar, UserGate и RuSIEM отдельно акцентируют масштабирование и/или кластерность.",
        "add": "Нужен production-grade scale-out для ClickHouse или эквивалент, replica/shard topology, failover и capacity planning как штатная функция.",
    },
    {
        "criterion": "Контент и detection coverage",
        "ours": "Базовый content layer уже есть: normalizer, filter, correlation packs, vuln/TI/source coverage. Но библиотека правил, парсеров и контент-паков существенно уже enterprise baseline.",
        "global": "Мировые лидеры сильны готовыми detection packs, content hub/marketplace, risk-based detections, UEBA/XDR content и регулярным vendor-updated контентом.",
        "russian": "Российские лидеры сильны локальными rule packs, отраслевыми сценариями и лучшей привязкой к местному threat landscape и типовым источникам российского enterprise-рынка.",
        "add": "Нужны versioned content bundles, CI/CD для контента, тестовые датасеты, release rings, richer parser catalog и заметное расширение rule packs.",
    },
    {
        "criterion": "UEBA / entity analytics / graph",
        "ours": "Есть entities, risk signals, humanization и AI assessment, но нет зрелой поведенческой модели пользователей, хостов и сервисов в масштабе enterprise UEBA.",
        "global": "Sentinel, Google SecOps, XSIAM, QRadar, Securonix, Exabeam и Rapid7 явно продвигают поведенческую аналитику, entity context, TDIR/AI и более зрелое поведенческое расследование.",
        "russian": "В РФ глубина UEBA различается, но рынок уже уходит от голой корреляции к richer context, AI-ассистентам и поведению сущностей.",
        "add": "Нужны baselines по user/host/service, граф сущностей и цепочек атаки, session/process lineage и cross-host behavior analytics.",
    },
    {
        "criterion": "SOAR и response automation",
        "ours": "Есть response actions, approvals, Telegram bot, safe host actions и базовая оркестрация. Это уже operational SOAR, но ещё не широкий enterprise action library.",
        "global": "Splunk SOAR, Cortex XSOAR/XSIAM, Sentinel playbooks, QRadar SOAR, Exabeam SOAR и часть других мировых решений дают более зрелую библиотеку playbooks и action integrations.",
        "russian": "R-Vision, Security Vision, Solar и UserGate особенно сильны в IRP/SOAR workflow, согласованиях, тикетинге и локальной интеграции с enterprise-процессами.",
        "add": "Нужны playbook library, connector SDK для response actions, ticketing/CMDB/IAM/EDR integrations и richer closed-loop remediation.",
    },
    {
        "criterion": "Интеграции и коннекторы",
        "ours": "Есть connectors, ingest APIs, source discovery, Windows/Linux/network/app coverage, но каталог интеграций пока ограничен и largely custom.",
        "global": "У мировых лидеров очень широкий и регулярно поддерживаемый каталог интеграций, часто с официальной сертификацией, managed onboarding и data source templates.",
        "russian": "Российские платформы сильны интеграцией с локальными СЗИ, ИТ-системами, экосистемой вендора и регуляторным контуром.",
        "add": "Нужен системный connector program: облака, SaaS, EDR, IDP, email, proxy, firewalls, NDR, K8s, CI/CD и identity/cloud posture data.",
    },
    {
        "criterion": "Vuln / TI / exposure context",
        "ours": "Есть vuln workflow, threat intel и AI-enriched investigations. Это сильнее среднего самописного SIEM, но пока без глубокой exposure graph/attack path модели.",
        "global": "Google SecOps, XSIAM, Rapid7, QRadar и часть других мировых лидеров заметно глубже связывают SIEM с TI, NDR/XDR, exposure и prioritization контекстом.",
        "russian": "Российские лидеры сильны там, где TI, реагирование и расследование увязаны с реальными SOC-процессами, регуляторным контуром и внедренческой экспертизой.",
        "add": "Нужна нормальная exposure graph модель: asset-service-account-identity-vuln-internet path, prioritization by blast radius и attack path view.",
    },
    {
        "criterion": "IAM / governance / secrets / audit",
        "ours": "Здесь у нас уже сильная база: OIDC-first, Vault, service accounts, token rotation, audit chain, control plane и access governance.",
        "global": "Enterprise лидеры обычно дополняют это зрелой federation-моделью, delegated admin, tenant admin boundaries и packaged policy stacks.",
        "russian": "Локальные лидеры часто хорошо покрывают корпоративные workflows, согласования, аудит и нормативную трассируемость.",
        "add": "Нужно добить SAML/LDAP/AD federation, delegated admin, stronger policy packs и tenant-level governance boundaries.",
    },
    {
        "criterion": "Multi-tenancy и MSSP readiness",
        "ours": "На текущем стенде это пока ограничено. Есть RBAC/access model, но нет зрелой tenant isolation и MSSP operating model.",
        "global": "Для мировых лидеров multi-workspace, tenant segregation и service-provider scenarios давно являются штатным enterprise-требованием.",
        "russian": "Solar прямо позиционируется и для MSSP-провайдеров, ArcSight historically силён в multi-tenant сценариях, а RuSIEM и UserGate явно говорят про масштабирование и сервисную модель.",
        "add": "Нужны tenant-scoped data/view/control planes, tenant-safe content promotion, billing/usage segmentation и delegated tenant administration.",
    },
    {
        "criterion": "UX и analyst workflow",
        "ours": "UI уже рабочий и русифицированный, есть incidents, events, assets, entities, builders, docs, access. Но уровню polished enterprise console пока мешают неоднородность и незавершённость части потоков.",
        "global": "У лидеров сильнее единообразие рабочих сценариев, investigation graph, timelines, hunt-workflows и role-based ergonomics.",
        "russian": "Сильные локальные продукты часто делают упор на регламентный SOC-процесс, IR/approval workflows и операционную прозрачность для крупных корпоративных команд.",
        "add": "Нужны единая investigation surface, richer timelines/graphs, stronger triage-to-case flow и content/admin UX уровня production product.",
    },
    {
        "criterion": "AI и analyst assist",
        "ours": "Есть OpenClaw AI assessment, external search enrichment и Telegram-assisted workflow. Это сильная заготовка, но пока advisory layer, а не полноразмерный AI SOC co-pilot.",
        "global": "Google, Microsoft, Palo Alto, Rapid7, Securonix и другие мировые лидеры активно продвигают AI-assisted triage, summarization, hunting и investigation acceleration как встроенную часть платформы.",
        "russian": "В РФ AI-слой ещё неоднороден, но Solar и часть новых платформ уже продвигают AI-ассистентов и автоматизированную аналитику.",
        "add": "Нужны grounded AI workflows с evidence graph, explainability, content QA и operator-safe action suggestions на уровне playbooks.",
    },
    {
        "criterion": "Соответствие enterprise-class ожиданиям",
        "ours": "Платформа уже пригодна для серьёзной эксплуатации в controlled environment, но пока ближе к сильно развитому собственному SOC stack, чем к finished enterprise product line.",
        "global": "Лидеры продают не только движок, но и ecosystem: контент, support, integrations, services, certifications и repeatable rollout model.",
        "russian": "Локальные лидеры дополнительно выигрывают наличием локальной экспертизы, внедренческой сети и адаптацией под регуляторный и отраслевой контекст.",
        "add": "Нужны industrialization, packaged delivery, supportability, documentation maturity, certification/test harness и operating model для крупных внедрений.",
    },
]

SCORECARD_ROWS = [
    ("Архитектура и масштабирование", "Сильная modular база", "Нет настоящего scale-out", "Критично"),
    ("Контент и detection coverage", "Рабочее ядро контента", "Мало rule packs и parser coverage", "Критично"),
    ("UEBA / entity analytics", "Есть entities и AI assist", "Нет зрелой UEBA-модели", "Высокий"),
    ("SOAR и automation", "Operational SOAR уже есть", "Узкая library действий", "Высокий"),
    ("Интеграции", "Есть connectors и ingest APIs", "Каталог интеграций узкий", "Критично"),
    ("Vuln / TI / exposure", "Есть vuln и TI контур", "Нет exposure graph", "Высокий"),
    ("IAM / governance", "Здесь база уже сильная", "Не хватает federation depth", "Средний"),
    ("Multi-tenancy / MSSP", "Пока ограничено", "Нет tenant isolation", "Высокий"),
    ("UX / analyst workflow", "UI уже рабочий", "Не хватает polished investigation UX", "Высокий"),
    ("AI / analyst assist", "Есть OpenClaw AI", "Пока advisory, не co-pilot", "Средний"),
]

WHAT_WE_HAVE = [
    "Прозрачная модульная архитектура с реальным разделением ingest, processing, storage и web/control planes.",
    "Kafka + ClickHouse pipeline с hot/cold retention и отдельным stream/batch detection path.",
    "Инциденты, кейсы, сущности, активы, connectors, access, docs, builders и русифицированный shell.",
    "SOAR-база: response actions, approvals, Telegram, safe host actions.",
    "Vulnerability и threat-intel контур, OpenClaw AI assessment, Windows/Linux telemetry expansion.",
    "OIDC-first, Vault, service accounts, rotation, audit chain и governance surfaces.",
]

WHAT_MISSING = [
    "Настоящее горизонтальное масштабирование и кластеризация data/search layer.",
    "Большой каталог поддерживаемых интеграций и контент-маркетплейс.",
    "Глубокая UEBA/entity analytics и attack graph/lineage модель.",
    "Полноценная multi-tenancy/MSSP-модель.",
    "Богатые compliance/regulatory packs и отраслевые сценарии.",
    "Глубокая XDR/EDR/NDR/identity/cloud fusion-модель.",
    "Более продуктовый admin/content lifecycle и polished analyst UX.",
]

ROADMAP = [
    ("Приоритет 0-3 месяца", [
        "Перевести storage/search layer в реальный scale-out режим: replicated/sharded ClickHouse или эквивалентный дизайн.",
        "Сделать industrialized content operations: versioning, signing, tests, promotion rings, rollback.",
        "Сильно расширить коннекторы: AD/Entra ID, VPN, firewalls, EDR, email, proxy, SaaS, cloud, Kubernetes, CI/CD.",
        "Добавить полноценный evidence graph: actor IP, user, process, host, service, account, session, parent-child lineage.",
        "Нормализовать retention/tiering как продуктовую настройку с долгосрочным архивом и быстрым rehydrate.",
    ]),
    ("Приоритет 3-6 месяцев", [
        "Ввести UEBA/baselines по user/host/service и risk-driven fusion детектов.",
        "Реализовать tenant isolation и MSSP-ready operating model.",
        "Сделать library playbooks и response connectors для IR/ITSM/IAM/EDR.",
        "Добавить compliance content packs и executive reporting.",
    ]),
    ("Приоритет 6-12 месяцев", [
        "Построить exposure graph и attack-path prioritization.",
        "Сделать зрелую DR/failover automation и cross-site standby.",
        "Собрать продуктовую delivery-модель: packaging, docs, upgrade tracks, validation kits, support runbooks.",
    ]),
]

VENDOR_NOTES = {
    "Splunk Enterprise Security / SOAR": ["Splunk ES остаётся эталоном по risk-based alerting, investigation workflow, content operations и зрелости экосистемы SOAR/integrations."],
    "Microsoft Sentinel": ["Sentinel — сильный baseline для cloud-native SIEM/SOAR, data connectors, automation/playbooks, content hub и UEBA/entity behavior analytics."],
    "Google Security Operations": ["Google SecOps силён связкой SIEM + TI + Mandiant + AI-assisted workflows и является ориентиром для TI-centric investigation."],
    "Palo Alto Cortex XSIAM": ["XSIAM — эталон AI-driven SOC платформы с сильной автоматизацией и интеграцией XDR/SIEM/SOAR."],
    "IBM QRadar": ["QRadar остаётся одним из классических enterprise-лидеров: сильный UBA/NDR/SOAR story, большой каталог интеграций и зрелый on-prem operating model."],
    "OpenText ArcSight": ["ArcSight по-прежнему важен как heavyweight SIEM с сильной real-time correlation, multi-tenancy, TI/SOAR и high-EPS позиционированием."],
    "Rapid7 InsightIDR / Incident Command": ["Rapid7 важен как сильный cloud-native SIEM/XDR ориентир с attack surface context, AI-driven triage, UEBA, TI и SOAR."],
    "Securonix Unified Defense SIEM": ["Securonix — заметный лидер next-gen SIEM с акцентом на Unified Defense SIEM, UEBA, SOAR, hot data platform и AI SOC analyst workflows."],
    "Exabeam Security Operations Platform": ["Exabeam важен как сильный UEBA/behavior analytics/SOAR игрок, особенно как ориентир по user timeline и поведению сущностей."],
    "Elastic Security / Elastic SIEM": ["Elastic важен как мощный search/analytics-based security platform пример с сильной масштабируемостью и гибкостью."],
    "MaxPatrol SIEM": ["MaxPatrol SIEM — один из ключевых ориентиров по российскому high-load SIEM, локальному контенту и detection экспертизе."],
    "Kaspersky SIEM / KUMA": ["KUMA важен как enterprise on-prem SIEM с multitenancy, категоризацией активов, интеграциями и структурированной локальной экосистемой."],
    "R-Vision SIEM / SOAR": ["R-Vision — сильный ориентир по IRP/SOAR процессам, согласованиям и SOC automation в российском контексте."],
    "Security Vision": ["Security Vision важен как platform-style стек SIEM/SOAR/IRP/GRC с упором на процессность и enterprise governance."],
    "UserGate SIEM": ["UserGate SIEM важен как быстро усиливающийся экосистемный игрок: SIEM + IRP/SOAR + TI, cluster architecture, EDR/NAC/VPN телеметрия и фокус на российский рынок."],
    "Solar SIEM": ["Solar SIEM важен как новый сильный локальный ориентир: единое ядро SIEM+SOAR, контент Solar JSOC, AI-ассистент и явный MSSP/мультизаказчиковый сценарий."],
    "RuSIEM": ["RuSIEM важен как отечественный игрок с real-time correlation, baseline-поведенческим анализом, масштабируемой микросервисной архитектурой и встроенным incident management."],
}

SOURCES = [
    ("Splunk Enterprise Security", "https://www.splunk.com/en_us/products/enterprise-security.html"),
    ("Splunk Enterprise Security Features", "https://www.splunk.com/en_us/products/splunk-enterprise-security-features.html"),
    ("Microsoft Sentinel overview", "https://learn.microsoft.com/en-us/azure/sentinel/overview"),
    ("Microsoft Sentinel data connectors", "https://learn.microsoft.com/en-us/azure/sentinel/data-connectors-reference"),
    ("Microsoft Sentinel automation", "https://learn.microsoft.com/en-us/azure/sentinel/automation"),
    ("Microsoft Sentinel UEBA", "https://learn.microsoft.com/en-us/azure/sentinel/identify-threats-with-entity-behavior-analytics"),
    ("Google Security Operations", "https://cloud.google.com/security/products/security-operations"),
    ("Palo Alto Cortex XSIAM", "https://www.paloaltonetworks.com/cortex/xsiam"),
    ("Palo Alto Cortex XSOAR", "https://www.paloaltonetworks.com/cortex/xsoar"),
    ("IBM QRadar SIEM", "https://www.ibm.com/products/qradar-siem"),
    ("IBM QRadar suite", "https://www.ibm.com/products/qradar"),
    ("OpenText ArcSight Enterprise Security Manager", "https://www.opentext.com/products/enterprise-security-manager"),
    ("Rapid7 InsightIDR", "https://www.rapid7.com/products/insightidr/"),
    ("Rapid7 InsightIDR docs", "https://docs.rapid7.com/insightidr/"),
    ("Securonix Unified Defense SIEM", "https://www.securonix.com/resources/unified-defense-siem/"),
    ("Securonix UEBA", "https://www.securonix.com/products/ueba/"),
    ("Exabeam SOAR", "https://www.exabeam.com/capabilities/soar/"),
    ("Elastic SIEM", "https://www.elastic.co/security/siem"),
    ("Positive Technologies MaxPatrol SIEM", "https://www.ptsecurity.com/ru-ru/products/mpsiem/"),
    ("Positive Technologies MaxPatrol SIEM 8.0", "https://ptsecurity.com/ru-ru/about/news/maxpatrol-siem-8-0-ml-dlya-povedencheskogo-analiza-snizhenie-apparatnyh-trebovanij-i-obrabotka-svyshe-500-tys-sobytij-v-sekundu/"),
    ("Kaspersky KUMA documentation", "https://support.kaspersky.com/KUMA/4.2/en-US/294293.htm"),
    ("Kaspersky Unified Monitoring and Analysis Platform", "https://www.kaspersky.com/enterprise-security/unified-monitoring-and-analysis-platform"),
    ("R-Vision SIEM", "https://rvision.ru/products/siem"),
    ("R-Vision SOAR", "https://rvision.ru/products/soar"),
    ("Security Vision products", "https://www.securityvision.ru/products/"),
    ("Security Vision SIEM", "https://www.securityvision.ru/products/siem/"),
    ("Security Vision SOAR", "https://www.securityvision.ru/products/soar/"),
    ("UserGate SIEM", "https://siem.usergate.com/"),
    ("UserGate SIEM architecture", "https://docs.usergate.com/arhitektura-siem_1570.html"),
    ("Solar SIEM", "https://rt-solar.ru/products/solar_siem/"),
    ("Solar SIEM architecture", "https://rt-solar.ru/products/solar_siem/architecture/"),
    ("RuSIEM", "https://rusiem.com/ru/products/rusiem"),
    ("Local architecture reference", str(ROOT / "docs" / "architecture.md")),
    ("Local enterprise foundation reference", str(ROOT / "docs" / "enterprise_foundation.md")),
    ("Local AI/SOAR reference", str(ROOT / "docs" / "openclaw_incident_ai_telegram_wave_2026-03-29.md")),
    ("Local EPS assessment", str(ROOT / "docs" / "performance_eps_assessment_2026-03-26.md")),
    ("Local governance closure", str(ROOT / "docs" / "production_certification_and_governance_closure_2026-03-26.md")),
    ("Local retention/storage hardening", str(ROOT / "docs" / "storage_rebalance_and_retention_hardening_2026-04-05.md")),
]


def write_markdown() -> None:
    lines = [
        f"# {TITLE}",
        "",
        SUBTITLE,
        "",
        "## Методика",
        "",
        "Сравнение основано на публичных официальных материалах поставщиков, просмотренных 8 апреля 2026 года, и на внутренних документах текущей платформы Rdegon SIEM. Эта версия расширяет baseline и явно включает IBM QRadar, ArcSight, Rapid7, Securonix, Exabeam, Elastic, UserGate, Solar и RuSIEM.",
        "",
        "## Executive Summary",
        "",
    ]
    lines.extend(f"- {item}" for item in EXEC_SUMMARY)
    lines.extend(["", "## Репрезентативный набор платформ", ""])
    for scope, vendors in MARKET_SCOPE.items():
        lines.append(f"### {scope}")
        lines.append("")
        lines.extend(f"- {vendor}" for vendor in vendors)
        lines.append("")
    lines.extend(["## Краткая scorecard", ""])
    lines.append("| Критерий | Наша позиция | Главный разрыв | Приоритет |")
    lines.append("| --- | --- | --- | --- |")
    for row in SCORECARD_ROWS:
        vals = [item.replace("|", "\\|") for item in row]
        lines.append("| " + " | ".join(vals) + " |")
    lines.extend(["", "## Подробный разбор по критериям", ""])
    for row in COMPARISON_ROWS:
        lines.append(f"### {row['criterion']}")
        lines.append("")
        lines.append(f"- Наша система: {row['ours']}")
        lines.append(f"- Мировые лидеры: {row['global']}")
        lines.append(f"- Российские лидеры: {row['russian']}")
        lines.append(f"- Что нам добавить: {row['add']}")
        lines.append("")
    lines.extend(["## Что у нас уже есть", ""])
    lines.extend(f"- {item}" for item in WHAT_WE_HAVE)
    lines.extend(["", "## Чего пока нет", ""])
    lines.extend(f"- {item}" for item in WHAT_MISSING)
    lines.extend(["", "## Вендорные ориентиры", ""])
    for vendor, notes in VENDOR_NOTES.items():
        lines.append(f"### {vendor}")
        lines.append("")
        lines.extend(f"- {note}" for note in notes)
        lines.append("")
    lines.extend(["## Рекомендуемый roadmap", ""])
    for phase, items in ROADMAP:
        lines.append(f"### {phase}")
        lines.append("")
        lines.extend(f"- {item}" for item in items)
        lines.append("")
    lines.extend(["## Источники", ""])
    lines.extend(f"- {label}: {url}" for label, url in SOURCES)
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(font_size)


def set_column_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def write_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    r.italic = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(10)
    doc.add_heading("Методика", level=1)
    doc.add_paragraph("Сравнение основано на публичных официальных материалах поставщиков, просмотренных 8 апреля 2026 года, и на внутренних документах текущей платформы Rdegon SIEM. Эта версия расширяет baseline и явно включает IBM QRadar, ArcSight, Rapid7, Securonix, Exabeam, Elastic, UserGate, Solar и RuSIEM.")
    doc.add_heading("Executive Summary", level=1)
    for item in EXEC_SUMMARY:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Репрезентативный набор платформ", level=1)
    for scope, vendors in MARKET_SCOPE.items():
        doc.add_heading(scope, level=2)
        for vendor in vendors:
            doc.add_paragraph(vendor, style="List Bullet")
    doc.add_heading("Краткая scorecard", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_column_widths(table, [5.0, 7.0, 10.5, 3.0])
    headers = ["Критерий", "Наша позиция", "Главный разрыв", "Приоритет"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, font_size=9)
    for row in SCORECARD_ROWS:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0], bold=True, font_size=8)
        set_cell_text(cells[1], row[1], font_size=8)
        set_cell_text(cells[2], row[2], font_size=8)
        set_cell_text(cells[3], row[3], font_size=8)
    doc.add_heading("Подробный разбор по критериям", level=1)
    for row in COMPARISON_ROWS:
        doc.add_heading(row["criterion"], level=2)
        doc.add_paragraph(f"Наша система: {row['ours']}", style="List Bullet")
        doc.add_paragraph(f"Мировые лидеры: {row['global']}", style="List Bullet")
        doc.add_paragraph(f"Российские лидеры: {row['russian']}", style="List Bullet")
        doc.add_paragraph(f"Что нам добавить: {row['add']}", style="List Bullet")
    for title, items in [("Что у нас уже есть", WHAT_WE_HAVE), ("Чего пока нет", WHAT_MISSING)]:
        doc.add_heading(title, level=1)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Вендорные ориентиры", level=1)
    for vendor, notes in VENDOR_NOTES.items():
        doc.add_heading(vendor, level=2)
        for note in notes:
            doc.add_paragraph(note, style="List Bullet")
    doc.add_heading("Рекомендуемый roadmap", level=1)
    for phase, items in ROADMAP:
        doc.add_heading(phase, level=2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Источники", level=1)
    for label, url in SOURCES:
        doc.add_paragraph(f"{label}: {url}", style="List Bullet")
    doc.save(DOCX_PATH)


def write_pdf() -> None:
    pdfmetrics.registerFont(TTFont("ArialCustom", r"C:\Windows\Fonts\arial.ttf"))
    pdfmetrics.registerFont(TTFont("ArialCustom-Bold", r"C:\Windows\Fonts\arialbd.ttf"))
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="RusTitle", parent=styles["Title"], fontName="ArialCustom-Bold", fontSize=18, leading=22, alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="RusSubtitle", parent=styles["Normal"], fontName="ArialCustom", fontSize=10, leading=12, alignment=TA_CENTER, textColor=colors.HexColor("#444444")))
    styles.add(ParagraphStyle(name="RusH1", parent=styles["Heading1"], fontName="ArialCustom-Bold", fontSize=14, leading=18, spaceAfter=8))
    styles.add(ParagraphStyle(name="RusH2", parent=styles["Heading2"], fontName="ArialCustom-Bold", fontSize=11, leading=14, spaceAfter=6))
    styles.add(ParagraphStyle(name="RusBody", parent=styles["BodyText"], fontName="ArialCustom", fontSize=9, leading=12))
    styles.add(ParagraphStyle(name="RusBullet", parent=styles["BodyText"], fontName="ArialCustom", fontSize=9, leading=12, leftIndent=12, bulletIndent=0))
    styles.add(ParagraphStyle(name="RusSmall", parent=styles["BodyText"], fontName="ArialCustom", fontSize=8, leading=10))
    story = [Paragraph(TITLE, styles["RusTitle"]), Spacer(1, 0.15 * cm), Paragraph(SUBTITLE, styles["RusSubtitle"]), Spacer(1, 0.4 * cm), Paragraph("Методика", styles["RusH1"]), Paragraph("Сравнение основано на публичных официальных материалах поставщиков, просмотренных 8 апреля 2026 года, и на внутренних документах текущей платформы Rdegon SIEM. Эта версия расширяет baseline и явно включает IBM QRadar, ArcSight, Rapid7, Securonix, Exabeam, Elastic, UserGate, Solar и RuSIEM.", styles["RusBody"]), Spacer(1, 0.2 * cm), Paragraph("Executive Summary", styles["RusH1"])]
    story.extend(Paragraph(f"• {item}", styles["RusBullet"]) for item in EXEC_SUMMARY)
    story.extend([Spacer(1, 0.2 * cm), Paragraph("Репрезентативный набор платформ", styles["RusH1"])])
    for scope, vendors in MARKET_SCOPE.items():
        story.append(Paragraph(scope, styles["RusH2"]))
        story.extend(Paragraph(f"• {vendor}", styles["RusBullet"]) for vendor in vendors)
    story.extend([Spacer(1, 0.2 * cm), Paragraph("Краткая scorecard", styles["RusH1"])])
    table_data = [["Критерий", "Наша позиция", "Главный разрыв", "Приоритет"]]
    for row in SCORECARD_ROWS:
        table_data.append(list(row))
    table = Table(table_data, colWidths=[5.0 * cm, 7.0 * cm, 12.0 * cm, 3.0 * cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
        ("FONTNAME", (0, 0), (-1, 0), "ArialCustom-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "ArialCustom"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.0),
        ("LEADING", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#808080")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FC")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("Подробный разбор по критериям", styles["RusH1"]))
    for row in COMPARISON_ROWS:
        story.append(Paragraph(row["criterion"], styles["RusH2"]))
        story.append(Paragraph(f"• Наша система: {row['ours']}", styles["RusBullet"]))
        story.append(Paragraph(f"• Мировые лидеры: {row['global']}", styles["RusBullet"]))
        story.append(Paragraph(f"• Российские лидеры: {row['russian']}", styles["RusBullet"]))
        story.append(Paragraph(f"• Что нам добавить: {row['add']}", styles["RusBullet"]))
    story.append(PageBreak())
    for title, items in [("Что у нас уже есть", WHAT_WE_HAVE), ("Чего пока нет", WHAT_MISSING)]:
        story.append(Paragraph(title, styles["RusH1"]))
        story.extend(Paragraph(f"• {item}", styles["RusBullet"]) for item in items)
        story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("Вендорные ориентиры", styles["RusH1"]))
    for vendor, notes in VENDOR_NOTES.items():
        story.append(Paragraph(vendor, styles["RusH2"]))
        story.extend(Paragraph(f"• {note}", styles["RusBullet"]) for note in notes)
    story.append(Paragraph("Рекомендуемый roadmap", styles["RusH1"]))
    for phase, items in ROADMAP:
        story.append(Paragraph(phase, styles["RusH2"]))
        story.extend(Paragraph(f"• {item}", styles["RusBullet"]) for item in items)
    story.append(PageBreak())
    story.append(Paragraph("Источники", styles["RusH1"]))
    for label, url in SOURCES:
        safe = f"{label}: {url}".replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, styles["RusSmall"]))
    pdf = SimpleDocTemplate(str(PDF_PATH), pagesize=landscape(A4), leftMargin=1.2 * cm, rightMargin=1.2 * cm, topMargin=1.0 * cm, bottomMargin=1.0 * cm)
    pdf.build(story)


def main() -> None:
    write_markdown()
    write_docx()
    write_pdf()
    print(MD_PATH)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
